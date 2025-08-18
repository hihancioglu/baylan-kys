import io
import os
from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document
from storage import _s3, S3_BUCKET


def _detect_ext(data: bytes) -> str:
    """Best effort detection of a PDF or DOCX file from bytes."""
    if data.startswith(b"%PDF"):
        return ".pdf"
    if data.startswith(b"PK"):
        return ".docx"
    return ""


def extract_text(key_or_bytes: str | bytes) -> str:
    """Extract text from a PDF or DOCX document.

    Accepts either an S3/MinIO object key or raw file bytes. When given a key,
    the object is downloaded into memory before processing.
    """
    try:
        data: bytes
        ext = ""
        if isinstance(key_or_bytes, bytes):
            data = key_or_bytes
            ext = _detect_ext(data)
        else:
            if os.path.exists(key_or_bytes):
                ext = os.path.splitext(key_or_bytes)[1].lower()
                if ext == ".pdf":
                    return pdf_extract_text(key_or_bytes)
                if ext == ".docx":
                    doc = Document(key_or_bytes)
                    return "\n".join(p.text for p in doc.paragraphs)
                return ""
            obj = _s3.get_object(Bucket=S3_BUCKET, Key=key_or_bytes)
            data = obj["Body"].read()
            ext = os.path.splitext(key_or_bytes)[1].lower()

        if ext == ".pdf":
            return pdf_extract_text(io.BytesIO(data))
        if ext == ".docx":
            doc = Document(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs)
        return ""
    except Exception:
        return ""
