import io
import os
from pathlib import Path
import sys
from unittest.mock import MagicMock

import pytest
from docx import Document as DocxDocument

# Ensure env vars for storage before importing app
os.environ.setdefault("S3_ENDPOINT", "http://s3")
os.environ.setdefault("S3_BUCKET_MAIN", "local")

repo_root = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(repo_root))
sys.path.insert(0, str(repo_root / "portal"))


@pytest.fixture()
def app_models():
    import app as app_module
    import models as models_module
    app_module.app.config["WTF_CSRF_ENABLED"] = False
    return app_module, models_module


@pytest.fixture()
def client(app_models):
    app_module, _ = app_models
    client = app_module.app.test_client()
    with client.session_transaction() as sess:
        sess["user"] = {"id": 1}
        sess["roles"] = ["reader"]
    return client


def _setup_document(models):
    session = models.SessionLocal()
    doc = models.Document(
        file_key="doc/latest",
        title="Doc",
        status="Published",
        major_version=1,
        minor_version=1,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    session.add(doc)
    session.commit()
    rev1 = models.DocumentRevision(
        doc_id=doc.id,
        major_version=1,
        minor_version=0,
        file_key="doc/v1",
    )
    rev2 = models.DocumentRevision(
        doc_id=doc.id,
        major_version=1,
        minor_version=1,
        file_key="doc/v2",
    )
    session.add_all([rev1, rev2])
    session.commit()
    doc_id, r1_id, r2_id = doc.id, rev1.id, rev2.id
    session.close()
    return doc_id, r1_id, r2_id


def _docx(text: str) -> bytes:
    buf = io.BytesIO()
    d = DocxDocument()
    d.add_paragraph(text)
    d.save(buf)
    return buf.getvalue()


def test_compare_api_creates_diff_and_returns_url(app_models, client):
    app_module, models = app_models
    doc_id, rev1_id, rev2_id = _setup_document(models)
    file1 = _docx("hello")
    file2 = _docx("world")

    app_module.storage_client.get_object = MagicMock(
        side_effect=[{"Body": io.BytesIO(file1)}, {"Body": io.BytesIO(file2)}]
    )
    app_module.storage_client.put = MagicMock()
    app_module.storage_client.generate_presigned_url = MagicMock(return_value="/signed")

    resp = client.post(f"/api/documents/{doc_id}/compare?from={rev1_id}&to={rev2_id}")
    assert resp.status_code == 200
    data = resp.get_json()
    expected = f"previews/{doc_id}/diff-{rev1_id}-{rev2_id}.html"
    assert data["filename"] == expected
    assert data["url"] == "/signed"
    app_module.storage_client.put.assert_called_once()
    app_module.storage_client.generate_presigned_url.assert_called_once_with(
        expected, bucket=app_module.storage_client.bucket_previews
    )


def test_compare_api_invalid_inputs_return_400(app_models, client):
    app_module, models = app_models
    doc_id, rev1_id, rev2_id = _setup_document(models)
    cases = [
        f"/api/documents/{doc_id}/compare",
        f"/api/documents/{doc_id}/compare?from={rev1_id}",
        f"/api/documents/{doc_id}/compare?from={rev1_id}&to={rev1_id}",
        f"/api/documents/{doc_id}/compare?from=abc&to={rev2_id}",
    ]
    for url in cases:
        resp = client.post(url)
        assert resp.status_code == 400


def test_compare_api_missing_revision_returns_404(app_models, client):
    app_module, models = app_models
    doc_id, rev1_id, rev2_id = _setup_document(models)
    missing = rev2_id + 100
    app_module.storage_client.get_object = MagicMock()
    resp = client.post(f"/api/documents/{doc_id}/compare?from={rev1_id}&to={missing}")
    assert resp.status_code == 404
    app_module.storage_client.get_object.assert_not_called()
