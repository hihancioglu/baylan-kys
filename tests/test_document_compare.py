import io
import os
from pathlib import Path
import sys
from unittest.mock import MagicMock

import pytest
from docx import Document as DocxDocument
import openpyxl

os.environ.setdefault("S3_ENDPOINT", "http://s3")
os.environ.setdefault("S3_BUCKET_MAIN", "local")

repo_root = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(repo_root))
sys.path.insert(0, str(repo_root / "portal"))


@pytest.fixture()
def app_models():
    import app as app_module
    import models as models_module
    app_module.app.config["WTF_CSRF_ENABLED"] = False
    return app_module, models_module


@pytest.fixture()
def client(app_models):
    app_module, _ = app_models
    client = app_module.app.test_client()
    with client.session_transaction() as sess:
        sess["user"] = {"id": 1}
        sess["roles"] = ["reader"]
    return client


def _setup_document(models, mime):
    session = models.SessionLocal()
    doc = models.Document(
        file_key="doc/latest",
        title="Doc",
        status="Published",
        major_version=1,
        minor_version=1,
        mime=mime,
    )
    session.add(doc)
    session.commit()
    rev1 = models.DocumentRevision(
        doc_id=doc.id,
        major_version=1,
        minor_version=0,
        file_key="doc/v1",
    )
    rev2 = models.DocumentRevision(
        doc_id=doc.id,
        major_version=1,
        minor_version=1,
        file_key="doc/v2",
    )
    session.add_all([rev1, rev2])
    session.commit()
    doc_id, r1_id, r2_id = doc.id, rev1.id, rev2.id
    session.close()
    return doc_id, r1_id, r2_id


def test_compare_word_versions_cached(app_models, client):
    app_module, models = app_models
    doc_id, rev1_id, rev2_id = _setup_document(
        models,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    def _docx(text: str) -> bytes:
        f = io.BytesIO()
        d = DocxDocument()
        d.add_paragraph(text)
        d.save(f)
        return f.getvalue()

    file1 = _docx("hello")
    file2 = _docx("hello world")
    app_module.storage_client.get_object = MagicMock(
        side_effect=[{"Body": io.BytesIO(file1)}, {"Body": io.BytesIO(file2)}]
    )

    resp = client.get(
        f"/documents/{doc_id}/compare?rev_id={rev1_id}&rev_id={rev2_id}"
    )
    assert resp.status_code == 200
    assert app_module.storage_client.get_object.call_count == 2

    session = models.SessionLocal()
    rev1 = session.get(models.DocumentRevision, rev1_id)
    assert rev1.compare_result is not None
    session.close()

    app_module.storage_client.get_object.reset_mock()
    resp = client.get(
        f"/documents/{doc_id}/compare?rev_id={rev1_id}&rev_id={rev2_id}"
    )
    assert resp.status_code == 200
    app_module.storage_client.get_object.assert_not_called()


def test_compare_excel_versions_cached(app_models, client):
    app_module, models = app_models
    doc_id, rev1_id, rev2_id = _setup_document(
        models,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    def _xlsx(value: str) -> bytes:
        f = io.BytesIO()
        wb = openpyxl.Workbook()
        ws = wb.active
        ws["A1"] = value
        wb.save(f)
        return f.getvalue()

    file1 = _xlsx("A")
    file2 = _xlsx("B")
    app_module.storage_client.get_object = MagicMock(
        side_effect=[{"Body": io.BytesIO(file1)}, {"Body": io.BytesIO(file2)}]
    )

    resp = client.get(
        f"/documents/{doc_id}/compare?rev_id={rev1_id}&rev_id={rev2_id}"
    )
    assert resp.status_code == 200
    assert app_module.storage_client.get_object.call_count == 2

    session = models.SessionLocal()
    rev1 = session.get(models.DocumentRevision, rev1_id)
    assert rev1.compare_result is not None
    session.close()

    app_module.storage_client.get_object.reset_mock()
    resp = client.get(
        f"/documents/{doc_id}/compare?rev_id={rev1_id}&rev_id={rev2_id}"
    )
    assert resp.status_code == 200
    app_module.storage_client.get_object.assert_not_called()
